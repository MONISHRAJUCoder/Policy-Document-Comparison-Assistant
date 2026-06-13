"""
Document Parser - Extracts text from various file formats (PDF, DOCX, TXT)
"""

import os


def extract_text(filepath: str) -> str:
    """
    Extract plain text from a file based on its extension.
    Supports: .txt, .pdf, .docx
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        return _read_txt(filepath)
    elif ext == '.pdf':
        return _read_pdf(filepath)
    elif ext == '.docx':
        return _read_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _read_txt(filepath: str) -> str:
    """Read plain text file."""
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file with supported encodings.")


def _read_pdf(filepath: str) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        text_parts = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_parts.append(page.get_text())
        doc.close()
        return '\n'.join(text_parts)
    except ImportError:
        # Fallback to pdfplumber if available
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                return '\n'.join(
                    page.extract_text() or '' for page in pdf.pages
                )
        except ImportError:
            raise ImportError(
                "PDF parsing requires PyMuPDF or pdfplumber. "
                "Install with: pip install pymupdf"
            )


def _read_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return '\n'.join(paragraphs)
    except ImportError:
        raise ImportError(
            "DOCX parsing requires python-docx. "
            "Install with: pip install python-docx"
        )
